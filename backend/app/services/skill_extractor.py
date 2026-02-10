from docx import Document # For DOCX text extraction
import PyPDF2 # For PDF text extraction
import re
from typing import Dict, List, Optional
from pathlib import Path # For file path handling
from app.services.skill_extractor import SkillExtractor

class ResumeParser:
    """Service for parsing resumes from PDF and DOCX formats"""
    
    def __init__(self): 
        # Common section headers to look for (used to split resume into sections)
        self.education_keywords = [
            'education', 'academic', 'qualification', 'degree', 'university',
            'college', 'school', 'bachelor', 'master', 'phd', 'diploma', 'academy'
        ]

        self.experience_keywords = [
            'experience', 'employment', 'work history', 'professional experience',
            'career', 'positions', 'work experience', 'employment history'
        ]

        self.skills_keywords = [
            'skills', 'technical skills', 'core competencies', 'expertise',
            'technologies', 'proficiencies', 'tools', 'programming languages'
        ]

        self.certifications_keywords = [
            'certifications', 'certificates', 'courses', 'training', 'licenses',
            'professional development'
        ]

        self.projects_keywords = [
            'projects', 'portfolio', 'case studies', 'work samples', 'assignments'
        ]

        self.summary_keywords = [
            'summary', 'profile', 'objective', 'about me', 'personal statement'
        ]
        
        # Initialize advanced skill extractor
        self.skill_extractor = SkillExtractor()

    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            text = "" 
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                # reader represents the PDF file and allows us to access its pages and extract text 
                # Extract text from all pages, page is 0-indexed but we want to loop through all pages
                for page in reader.pages:
                    page_text = page.extract_text() 
                    if page_text:
                        text += page_text + "\n"
            
            return text.strip()
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs as well as tables (some resumes use tables for layout)
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            # Also extract text from tables
            for table in doc.tables: # tables is a list of all tables in the document, each table has rows and cells
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += "\n" + cell.text
            
            return text.strip()
        except Exception as e:
            raise Exception(f"Error reading DOCX: {str(e)}")
    
    def extract_contact_info(self, text: str) -> Dict:
        """Extract contact information from resume text"""
        contact_info = {}
        
        # Extract email, re for regular expression to find email addresses in the text, we look for patterns that match typical email formats
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        contact_info['email'] = emails[0] if emails else None
        
        # Extract phone numbers (multiple formats - Saudi Arabia patterns)
        phone_patterns = [
            r'\+966\s?5\d\s?\d{3}\s?\d{4}',     # +966 5X XXX XXXX
            r'00966\s?5\d\s?\d{3}\s?\d{4}',     # 00966 5X XXX XXXX
            r'05\d\s?\d{3}\s?\d{4}',            # 05X XXX XXXX
            r'5\d\s?\d{3}\s?\d{4}',             # 5X XXX XXXX (no leading 0)
            # US/International formats (fallback)
            r'\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\(\d{3}\)\s?\d{3}[-.\s]?\d{4}',
            r'\d{3}[-.\s]\d{3}[-.\s]\d{4}',
        ]
        
        for pattern in phone_patterns:
            phones = re.findall(pattern, text)
            if phones:
                contact_info['phone'] = phones[0]
                break
        
        if 'phone' not in contact_info:
            contact_info['phone'] = None
        
        # Extract LinkedIn
        linkedin_pattern = r'(?:https?://)?(?:www\.)?linkedin\.com/in/[\w\-]+'
        linkedin = re.findall(linkedin_pattern, text.lower())
        contact_info['linkedin'] = linkedin[0] if linkedin else None
        
        # Extract GitHub
        github_pattern = r'(?:https?://)?(?:www\.)?github\.com/[\w\-]+'
        github = re.findall(github_pattern, text.lower())
        contact_info['github'] = github[0] if github else None
        
        # Extract name (usually first line or near top)
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        if lines:
            # First non-empty line is often the name
            potential_name = lines[0]
            # Check if it looks like a name (2-4 words, no special chars)
            if 2 <= len(potential_name.split()) <= 4 and re.match(r'^[A-Za-z\s]+$', potential_name):
                contact_info['name'] = potential_name
            else:
                contact_info['name'] = None
        else:
            contact_info['name'] = None
        
        return contact_info
    
    def extract_sections(self, text: str) -> Dict[str, str]:
        """Split resume into sections based on headers"""
        sections = { # Initialize all possible sections as empty strings
            'summary': '',
            'education': '',
            'experience': '',
            'skills': '',
            'certifications': '',
            'projects': '',
            'other': ''
        }
        
        lines = text.split('\n')
        current_section = 'other' # Default section if no headers found
        
        for line in lines:
            line_lower = line.lower().strip()
            
            # Check if line is a section header
            if any(keyword in line_lower for keyword in self.summary_keywords):
                current_section = 'summary'
                continue
            elif any(keyword in line_lower for keyword in self.education_keywords):
                current_section = 'education'
                continue
            elif any(keyword in line_lower for keyword in self.experience_keywords):
                current_section = 'experience'
                continue
            elif any(keyword in line_lower for keyword in self.skills_keywords):
                current_section = 'skills'
                continue
            elif any(keyword in line_lower for keyword in self.certifications_keywords):
                current_section = 'certifications'
                continue
            elif any(keyword in line_lower for keyword in self.projects_keywords):
                current_section = 'projects'
                continue
            
            # Add line to current section
            sections[current_section] += line + '\n'
        
        return sections
    
    def parse_education(self, education_text: str) -> List[Dict]:
        """Parse education section"""
        if not education_text.strip():
            return []
        
        education_entries = []
        
        # Common degree patterns
        degree_patterns = [
            r'(Bachelor|B\.S\.|B\.A\.|BS|BA|B\.Sc\.|B\.Tech)',
            r'(Master|M\.S\.|M\.A\.|MS|MA|M\.Sc\.|M\.Tech|MBA)',
            r'(Ph\.?D|Doctorate|Doctor)',
            r'(Associate|A\.S\.|A\.A\.)',
            r'(Diploma|Certificate)'
        ]
        
        lines = [line.strip() for line in education_text.split('\n') if line.strip()]
        
        current_entry = {}
        for line in lines:
            # Check for degree
            for pattern in degree_patterns:
                if re.search(pattern, line, re.IGNORECASE):
                    if current_entry:
                        education_entries.append(current_entry)
                    current_entry = {'degree': line}
                    break
            
            # Check for years (e.g., 2018-2022, 2020)
            year_match = re.search(r'\b(19|20)\d{2}\b', line)
            if year_match and 'year' not in current_entry:
                current_entry['year'] = year_match.group()
            
            # If line contains "University" or "College", it's likely the institution
            if any(word in line.lower() for word in ['university', 'college', 'institute', 'school']):
                if 'institution' not in current_entry:
                    current_entry['institution'] = line
        
        # Add last entry
        if current_entry:
            education_entries.append(current_entry)
        
        # If we didn't find structured entries, return the whole section as raw text
        if not education_entries and education_text.strip():
            education_entries.append({'raw': education_text.strip()})
        
        return education_entries
    
    def parse_experience(self, experience_text: str) -> List[Dict]:
        """Parse work experience section"""
        if not experience_text.strip():
            return []
        
        experience_entries = []
        lines = [line.strip() for line in experience_text.split('\n') if line.strip()]
        
        current_entry = {}
        for line in lines:
            # Check for date ranges (e.g., "Jan 2020 - Dec 2022" or "2020-2022")
            date_pattern = r'(\w{3,9}\s+\d{4}\s*[-–]\s*\w{3,9}\s+\d{4})|(\d{4}\s*[-–]\s*\d{4})|(\d{4}\s*[-–]\s*Present)'
            date_match = re.search(date_pattern, line, re.IGNORECASE)
            
            if date_match:
                if current_entry:
                    experience_entries.append(current_entry)
                current_entry = {'duration': date_match.group()}
                # The rest of the line might be the job title
                title = line.replace(date_match.group(), '').strip()
                if title:
                    current_entry['title'] = title
            
            # Check for company indicators
            elif any(indicator in line.lower() for indicator in ['inc.', 'llc', 'ltd', 'corporation', 'company']):
                if 'company' not in current_entry:
                    current_entry['company'] = line
            
            # Otherwise, it might be a description or bullet point
            elif current_entry and line:
                if 'description' not in current_entry:
                    current_entry['description'] = line
                else:
                    current_entry['description'] += '\n' + line
        
        # Add last entry
        if current_entry:
            experience_entries.append(current_entry)
        
        # If we didn't find structured entries, return the whole section
        if not experience_entries and experience_text.strip():
            experience_entries.append({'raw': experience_text.strip()})
        
        return experience_entries
    
    def parse_certifications(self, certifications_text: str) -> List[Dict]:
        """Parse certifications section"""
        if not certifications_text.strip():
            return []
        
        certifications_entries = []
        lines = [line.strip() for line in certifications_text.split('\n') if line.strip()]
        
        current_entry = {}
        for line in lines:
            # Check for year patterns
            year_match = re.search(r'\b(19|20)\d{2}\b', line)
            
            # If line has year, it's likely a certification with issue date
            if year_match:
                if current_entry:
                    certifications_entries.append(current_entry)
                current_entry = {
                    'name': line.replace(year_match.group(), '').strip(),
                    'year': year_match.group()
                }
            # Check for common certification providers
            elif any(provider in line.lower() for provider in ['aws', 'microsoft', 'google', 'cisco', 'comptia', 'oracle', 'pmi', 'certified']):
                if current_entry:
                    certifications_entries.append(current_entry)
                current_entry = {'name': line}
            # Add to existing entry or start new one
            elif line and not current_entry:
                current_entry = {'name': line}
            elif line and current_entry and 'description' not in current_entry:
                current_entry['description'] = line
        
        # Add last entry
        if current_entry:
            certifications_entries.append(current_entry)
        
        # If we didn't find structured entries, return the whole section
        if not certifications_entries and certifications_text.strip():
            certifications_entries.append({'raw': certifications_text.strip()})
        
        return certifications_entries
    
    def parse_projects(self, projects_text: str) -> List[Dict]:
        """Parse projects section"""
        if not projects_text.strip():
            return []
        
        projects_entries = []
        lines = [line.strip() for line in projects_text.split('\n') if line.strip()]
        
        current_entry = {}
        for line in lines:
            # Check for date ranges or years
            date_pattern = r'(\w{3,9}\s+\d{4}\s*[-–]\s*\w{3,9}\s+\d{4})|(\d{4}\s*[-–]\s*\d{4})|(\d{4}\s*[-–]\s*Present)|(\b(19|20)\d{2}\b)'
            date_match = re.search(date_pattern, line, re.IGNORECASE)
            
            # Check for project indicators (URLs, GitHub links, etc.)
            has_link = 'github.com' in line.lower() or 'http' in line.lower()
            
            # If line has date or looks like a project title
            if date_match or (len(line.split()) <= 8 and not current_entry):
                if current_entry:
                    projects_entries.append(current_entry)
                
                project_name = line
                if date_match:
                    project_name = line.replace(date_match.group(), '').strip()
                
                current_entry = {'name': project_name}
                if date_match:
                    current_entry['duration'] = date_match.group()
            
            # Check for GitHub or project links
            elif has_link:
                if 'link' not in current_entry:
                    current_entry['link'] = line
                elif 'description' not in current_entry:
                    current_entry['description'] = line
                else:
                    current_entry['description'] += '\n' + line
            
            # Add to description
            elif current_entry and line:
                if 'description' not in current_entry:
                    current_entry['description'] = line
                else:
                    current_entry['description'] += '\n' + line
            
            # Start new entry if no current entry
            elif not current_entry and line:
                current_entry = {'name': line}
        
        # Add last entry
        if current_entry:
            projects_entries.append(current_entry)
        
        # If we didn't find structured entries, return the whole section
        if not projects_entries and projects_text.strip():
            projects_entries.append({'raw': projects_text.strip()})
        
        return projects_entries
    
    def extract_skills(self, skills_text: str, full_text: str) -> List[Dict[str, str]]:
        """Extract skills using advanced NLP extractor"""
        # Combine skills section and full text for comprehensive extraction
        search_text = skills_text + '\n' + full_text
        
        # Use advanced skill extractor with NLP
        skills = self.skill_extractor.extract_skills(search_text)
        
        return skills
    
    def parse_resume(self, file_path: str, file_type: str) -> Dict:
        """
        Main method to parse resume
        Returns dict with all extracted information
        """
        # Extract raw text based on file type
        if file_type == 'pdf':
            raw_text = self.extract_text_from_pdf(file_path)
        elif file_type == 'docx':
            raw_text = self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}. Please use 'pdf' or 'docx' only!")
        
        # Extract contact information
        contact_info = self.extract_contact_info(raw_text)
        
        # Split into sections
        sections = self.extract_sections(raw_text)
        
        # Parse education
        education = self.parse_education(sections['education'])
        
        # Parse experience
        experience = self.parse_experience(sections['experience'])
        
        # Parse certifications
        certifications = self.parse_certifications(sections['certifications'])
        
        # Parse projects
        projects = self.parse_projects(sections['projects'])
        
        # Extract skills using advanced NLP extractor
        skills = self.extract_skills(sections['skills'], raw_text)
        
        return {
            'raw_text': raw_text,
            'contact_info': contact_info,
            'education': education,
            'experience': experience,
            'skills': skills,  # Now returns List[Dict] with 'name' and 'category'
            'certifications': certifications,
            'projects': projects,
            'sections': sections
        }