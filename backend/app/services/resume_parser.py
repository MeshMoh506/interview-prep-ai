from docx import Document
import PyPDF2
import re
from typing import Dict, List, Optional
from pathlib import Path

class ResumeParser:
    """Service for parsing resumes from PDF and DOCX formats"""
    
    def __init__(self):
        self.education_keywords = [
            'education', 'academic background', 'qualification'
        ]
        
        self.experience_keywords = [
            'cooperative training', 'work experience', 'professional experience',
            'employment', 'internship', 'experience', 'training'
        ]
        
        self.skills_keywords = [
            'technical skills', 'key skills', 'skills', 'competencies'
        ]
        
        self.projects_keywords = [
            'key projects', 'projects', 'portfolio'
        ]
        
        self.certifications_keywords = [
            'courses and certificates', 'certifications', 'courses', 'certificates'
        ]
        
        self.summary_keywords = [
            'summary', 'profile', 'objective'
        ]
        
        self.other_keywords = [
            'other', 'additional information'
        ]
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text.strip()
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += "\n" + cell.text
            
            return text.strip()
        except Exception as e:
            raise Exception(f"Error reading DOCX: {str(e)}")
    
    def extract_contact_info(self, text: str) -> Dict:
        """Extract contact information from resume text"""
        contact_info = {}
        
        # Extract email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        contact_info['email'] = emails[0] if emails else None
        
        # Extract phone
        phone_patterns = [
            r'\+966\s?5\d\s?\d{3}\s?\d{4}',
            r'00966\s?5\d\s?\d{3}\s?\d{4}',
            r'05\d\s?\d{3}\s?\d{4}',
            r'\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
        ]
        
        for pattern in phone_patterns:
            phones = re.findall(pattern, text)
            if phones:
                contact_info['phone'] = phones[0].strip()
                break
        
        if 'phone' not in contact_info:
            contact_info['phone'] = None
        
        # LinkedIn - Look for actual URL or username
        linkedin_pattern = r'(?:https?://)?(?:www\.)?linkedin\.com/in/([\w\-]+)'
        linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
        if linkedin_match:
            contact_info['linkedin'] = f'linkedin.com/in/{linkedin_match.group(1)}'
        else:
            # Check if "LinkedIn" text exists with a potential username nearby
            lines = text.split('\n')
            for line in lines[:5]:  # Check first 5 lines
                if 'linkedin' in line.lower() and '|' in line:
                    # Try to extract username from context
                    contact_info['linkedin'] = 'See resume header'
                    break
            if 'linkedin' not in contact_info:
                contact_info['linkedin'] = None
        
        # GitHub - same approach
        github_pattern = r'(?:https?://)?(?:www\.)?github\.com/([\w\-]+)'
        github_match = re.search(github_pattern, text, re.IGNORECASE)
        if github_match:
            contact_info['github'] = f'github.com/{github_match.group(1)}'
        else:
            lines = text.split('\n')
            for line in lines[:5]:
                if 'github' in line.lower() and '|' in line:
                    contact_info['github'] = 'See resume header'
                    break
            if 'github' not in contact_info:
                contact_info['github'] = None
        
        # Extract name (first line, 2-5 words)
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        if lines:
            potential_name = lines[0]
            if 2 <= len(potential_name.split()) <= 5 and re.match(r'^[A-Za-z\s\-]+$', potential_name):
                contact_info['name'] = potential_name
            else:
                contact_info['name'] = None
        else:
            contact_info['name'] = None
        
        return contact_info
    
    def extract_sections(self, text: str) -> Dict[str, str]:
        """Split resume into sections"""
        sections = {
            'summary': '',
            'education': '',
            'experience': '',
            'skills': '',
            'certifications': '',
            'projects': '',
            'other': ''
        }
        
        lines = text.split('\n')
        current_section = 'other'
        skip_next = False
        
        for i, line in enumerate(lines):
            if skip_next:
                skip_next = False
                continue
            
            line_lower = line.lower().strip()
            
            if not line_lower:
                continue
            
            # Detect section headers (ALL CAPS or short lines)
            is_potential_header = (
                line.strip().isupper() or 
                (len(line_lower.split()) <= 4 and not line.startswith('o'))
            )
            
            if is_potential_header:
                matched = False
                
                # Check each section type (order matters!)
                if any(kw in line_lower for kw in self.summary_keywords):
                    current_section = 'summary'
                    matched = True
                elif any(kw in line_lower for kw in self.experience_keywords):
                    current_section = 'experience'
                    matched = True
                elif any(kw in line_lower for kw in self.education_keywords):
                    current_section = 'education'
                    matched = True
                elif any(kw in line_lower for kw in self.projects_keywords):
                    current_section = 'projects'
                    matched = True
                elif any(kw in line_lower for kw in self.certifications_keywords):
                    current_section = 'certifications'
                    matched = True
                elif any(kw in line_lower for kw in self.skills_keywords):
                    current_section = 'skills'
                    matched = True
                elif any(kw in line_lower for kw in self.other_keywords):
                    current_section = 'other'
                    matched = True
                
                if matched:
                    continue
            
            # Add line to current section
            sections[current_section] += line + '\n'
        
        return sections
    
    def parse_education(self, education_text: str) -> List[Dict]:
        """Parse education section"""
        if not education_text.strip():
            return []
        
        education_entries = []
        lines = [line.strip() for line in education_text.split('\n') if line.strip()]
        
        current_entry = None
        
        for line in lines:
            # Institution line (has University, Academy, School, College)
            if any(word in line for word in ['University', 'Academy', 'School', 'College', 'Institute']):
                # Save previous
                if current_entry and 'institution' in current_entry:
                    education_entries.append(current_entry)
                
                current_entry = {'institution': line}
            
            # Degree line
            elif current_entry and any(word in line for word in ['Bachelor', 'Master', 'Degree', 'Diploma', 'Bootcamp', 'Certificate']):
                current_entry['degree'] = line
            
            # Year
            elif current_entry:
                year_match = re.search(r'\b(19|20)\d{2}\b', line)
                if year_match and 'year' not in current_entry:
                    current_entry['year'] = year_match.group()
                
                # Description (lines starting with 'o' or '-')
                if line.startswith('o') or line.startswith('-') or line.startswith('•'):
                    if 'description' not in current_entry:
                        current_entry['description'] = line
                    else:
                        current_entry['description'] += ' ' + line
        
        # Add last entry
        if current_entry and 'institution' in current_entry:
            education_entries.append(current_entry)
        
        return education_entries if education_entries else []
    
    def parse_experience(self, experience_text: str) -> List[Dict]:
        """Parse work experience section"""
        if not experience_text.strip():
            return []
        
        experience_entries = []
        lines = [line.strip() for line in experience_text.split('\n') if line.strip()]
        
        current_entry = None
        
        for line in lines:
            # Look for company/title line (has ● or —)
            if line.startswith('●') or '—' in line:
                # Save previous
                if current_entry and ('title' in current_entry or 'company' in current_entry):
                    experience_entries.append(current_entry)
                
                # Parse this line
                clean_line = line.lstrip('●').strip()
                
                if '—' in clean_line:
                    parts = clean_line.split('—')
                    current_entry = {
                        'company': parts[0].strip(),
                        'title': parts[1].strip() if len(parts) > 1 else ''
                    }
                else:
                    current_entry = {'company': clean_line}
            
            # Look for duration (2025/ 4 months) or (Jan 2020 - Dec 2022)
            elif current_entry:
                date_patterns = [
                    r'\d{4}/\s*\d+\s*months?',  # 2025/ 4 months
                    r'\d{4}\s*[-–]\s*\d{4}',    # 2020-2022
                    r'\d{4}\s*[-–]\s*Present',  # 2020-Present
                    r'\w{3,9}\s+\d{4}\s*[-–]\s*\w{3,9}\s+\d{4}'  # Jan 2020 - Dec 2022
                ]
                
                for pattern in date_patterns:
                    date_match = re.search(pattern, line, re.IGNORECASE)
                    if date_match and 'duration' not in current_entry:
                        current_entry['duration'] = date_match.group().strip()
                        break
                
                # Description (lines starting with 'o')
                if line.startswith('o') or line.startswith('-') or line.startswith('•'):
                    if 'description' not in current_entry:
                        current_entry['description'] = line
                    else:
                        current_entry['description'] += '\n' + line
        
        # Add last entry
        if current_entry and ('title' in current_entry or 'company' in current_entry):
            experience_entries.append(current_entry)
        
        return experience_entries
    
    def parse_projects(self, projects_text: str) -> List[Dict]:
        """Parse projects section"""
        if not projects_text.strip():
            return []
        
        projects = []
        lines = [line.strip() for line in projects_text.split('\n') if line.strip()]
        
        current_project = None
        
        for line in lines:
            # Project titles start with ●
            if line.startswith('●'):
                # Save previous
                if current_project and 'name' in current_project:
                    projects.append(current_project)
                
                project_name = line.lstrip('●').strip()
                current_project = {'name': project_name}
            
            # Descriptions start with 'o'
            elif current_project and (line.startswith('o') or line.startswith('-')):
                if 'description' not in current_project:
                    current_project['description'] = line
                else:
                    current_project['description'] += '\n' + line
        
        # Add last project
        if current_project and 'name' in current_project:
            projects.append(current_project)
        
        return projects
    
    def parse_certifications(self, certifications_text: str) -> List[Dict]:
        """Parse certifications/courses"""
        if not certifications_text.strip():
            return []
        
        certifications = []
        lines = [line.strip() for line in certifications_text.split('\n') if line.strip()]
        
        for line in lines:
            if line.startswith('o') or line.startswith('-') or line.startswith('•'):
                cert_name = line.lstrip('o-•').strip()
                
                # Extract year if present
                year_match = re.search(r'\b(19|20)\d{2}\b', cert_name)
                
                if year_match:
                    certifications.append({
                        'name': cert_name.replace(year_match.group(), '').strip(),
                        'year': year_match.group()
                    })
                else:
                    certifications.append({'name': cert_name})
        
        return certifications
    
    def extract_skills_from_text(self, skills_text: str, full_text: str) -> List[Dict]:
        """Extract skills from skills section"""
        skills_list = []
        seen_skills = set()
        
        # Comprehensive skills database
        skills_db = {
            'Programming Languages': ['Python', 'Java', 'JavaScript', 'TypeScript', 'C++', 'C#', 'C', 'Ruby', 'Go', 'Rust', 'Swift', 'Kotlin', 'PHP', 'Scala', 'R', 'MATLAB', 'Dart'],
            'Web Frontend': ['HTML', 'CSS', 'React', 'Angular', 'Vue', 'Vue.js', 'Next.js', 'Redux', 'Tailwind', 'Bootstrap', 'SASS', 'Webpack'],
            'Web Backend': ['Node.js', 'Express', 'Django', 'Flask', 'FastAPI', 'Fast API', 'Spring Boot', 'ASP.NET', 'Laravel', 'Rails'],
            'Mobile': ['Android', 'iOS', 'React Native', 'Flutter', 'Xamarin'],
            'Databases': ['SQL', 'PostgreSQL', 'MySQL', 'MongoDB', 'Redis', 'Elasticsearch', 'DynamoDB', 'Oracle', 'SQLite'],
            'Cloud & DevOps': ['AWS', 'Azure', 'GCP', 'Docker', 'Kubernetes', 'OpenShift', 'Jenkins', 'CI/CD', 'Terraform', 'Ansible', 'Linux', 'Ubuntu', 'Windows', 'Mac'],
            'Tools & Technologies': ['Git', 'GitHub', 'GitLab', 'Jira', 'Agile', 'Scrum', 'REST API', 'GraphQL', 'Microservices', 'UML', 'ERP', 'Odoo'],
            'Data & Analytics': ['Machine Learning', 'Deep Learning', 'TensorFlow', 'PyTorch', 'Pandas', 'NumPy', 'Data Analysis']
        }
        
        # Combine skills section and full text
        search_text = (skills_text + '\n' + full_text).lower()
        
        # Search for each skill
        for category, skills in skills_db.items():
            for skill in skills:
                skill_lower = skill.lower()
                
                # Use word boundary matching
                pattern = r'\b' + re.escape(skill_lower) + r'\b'
                
                if re.search(pattern, search_text) and skill not in seen_skills:
                    skills_list.append({
                        'name': skill,
                        'category': category
                    })
                    seen_skills.add(skill)
        
        return skills_list
    
    def parse_resume(self, file_path: str, file_type: str) -> Dict:
        """Main parse method"""
        # Extract text
        if file_type == 'pdf':
            raw_text = self.extract_text_from_pdf(file_path)
        elif file_type == 'docx':
            raw_text = self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        # Extract all sections
        contact_info = self.extract_contact_info(raw_text)
        sections = self.extract_sections(raw_text)
        
        education = self.parse_education(sections['education'])
        experience = self.parse_experience(sections['experience'])
        projects = self.parse_projects(sections['projects'])
        certifications = self.parse_certifications(sections['certifications'])
        skills = self.extract_skills_from_text(sections['skills'], raw_text)
        
        return {
            'raw_text': raw_text,
            'contact_info': contact_info,
            'education': education,
            'experience': experience,
            'skills': skills,
            'certifications': certifications,
            'projects': projects,
            'sections': sections
        }
