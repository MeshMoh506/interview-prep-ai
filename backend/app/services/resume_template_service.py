from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, List
import os

class ResumeTemplateService:
    """Generate professional resume documents from parsed data and selected template """
    # Define some basic templates with different styles and color schemes and can be easily extended in the future. Each template can have specific formatting rules
    TEMPLATES = {
        "professional": {
            "name": "Professional",
            "description": "Clean traditional format for corporate jobs",
            "colors": {"primary": (44, 62, 80), "secondary": (52, 152, 219)},
            "font": "Calibri",
            "suitable_for": ["Software Engineer", "Business Analyst", "Project Manager"]
        },
        "modern": {
            "name": "Modern",
            "description": "Contemporary design for tech roles",
            "colors": {"primary": (41, 128, 185), "secondary": (46, 204, 113)},
            "font": "Arial",
            "suitable_for": ["Frontend Developer", "Full Stack", "UI/UX Designer"]
        },
        "minimal": {
            "name": "Minimal",
            "description": "Simple clean design, ATS-friendly",
            "colors": {"primary": (0, 0, 0), "secondary": (100, 100, 100)},
            "font": "Times New Roman",
            "suitable_for": ["Any role - Maximum ATS compatibility"]
        }
    }
    
    def __init__(self):
        self.output_dir = "generated_resumes"
        os.makedirs(self.output_dir, exist_ok=True)
    
    def get_templates(self) -> List[Dict]:
        """Get all available templates"""
        return [
            {
                "id": key,
                "name": template["name"],
                "description": template["description"],
                "suitable_for": template["suitable_for"]
            }
            for key, template in self.TEMPLATES.items()
        ]
    
    def generate_resume(
        self,
        resume_data: Dict,
        template_id: str = "professional",
        user_id: int = 0
    ) -> Dict:
        """
        Generate a formatted resume DOCX file
        
        Args:
            resume_data: Parsed resume data from database
            template_id: Template to use (professional/modern/minimal)
            user_id: User ID for file naming
            
        Returns:
            Dict with file_path and success status
        """
        
        template = self.TEMPLATES.get(template_id, self.TEMPLATES["professional"])
        
        try:
            doc = Document()
            
            # Page setup
            self._setup_page(doc)
            
            # Build resume sections
            contact = resume_data.get("contact_info") or {}
            experience = resume_data.get("experience") or []
            education = resume_data.get("education") or []
            skills = resume_data.get("skills") or []
            projects = resume_data.get("projects") or []
            certifications = resume_data.get("certifications") or []
            
            # Add sections
            self._add_header(doc, contact, template)
            self._add_divider(doc, template)
            
            if resume_data.get("summary"):
                self._add_section(doc, "PROFESSIONAL SUMMARY", template)
                self._add_paragraph(doc, resume_data["summary"])
                self._add_divider(doc, template)
            
            if experience:
                self._add_section(doc, "WORK EXPERIENCE", template)
                self._add_experience(doc, experience, template)
                self._add_divider(doc, template)
            
            if education:
                self._add_section(doc, "EDUCATION", template)
                self._add_education(doc, education, template)
                self._add_divider(doc, template)
            
            if skills:
                self._add_section(doc, "TECHNICAL SKILLS", template)
                self._add_skills(doc, skills, template)
                self._add_divider(doc, template)
            
            if projects:
                self._add_section(doc, "PROJECTS", template)
                self._add_projects(doc, projects, template)
                self._add_divider(doc, template)
            
            if certifications:
                self._add_section(doc, "CERTIFICATIONS", template)
                self._add_certifications(doc, certifications, template)
            
            # Save file
            filename = f"resume_{user_id}_{template_id}.docx"
            filepath = os.path.join(self.output_dir, filename)
            doc.save(filepath)
            
            return {
                "success": True,
                "file_path": filepath,
                "filename": filename,
                "template_used": template["name"]
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Generation failed: {str(e)}"
            }
    
    def _setup_page(self, doc: Document):
        """Set page margins"""
        for section in doc.sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
    
    def _add_header(self, doc: Document, contact: Dict, template: Dict):
        """Add name and contact information"""
        primary = template["colors"]["primary"]
        secondary = template["colors"]["secondary"]
        font_name = template["font"]
        
        # Full Name
        name = contact.get("name", "Your Name")
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(name.upper())
        name_run.bold = True
        name_run.font.size = Pt(20)
        name_run.font.color.rgb = RGBColor(*primary)
        name_run.font.name = font_name
        
        # Contact line
        contact_parts = []
        if contact.get("email"):
            contact_parts.append(f"📧 {contact['email']}")
        if contact.get("phone"):
            contact_parts.append(f"📱 {contact['phone']}")
        if contact.get("linkedin"):
            contact_parts.append(f"💼 {contact['linkedin']}")
        if contact.get("github"):
            contact_parts.append(f"💻 {contact['github']}")
        
        if contact_parts:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_para.add_run("  |  ".join(contact_parts))
            contact_run.font.size = Pt(9)
            contact_run.font.color.rgb = RGBColor(*secondary)
            contact_run.font.name = font_name
    
    def _add_divider(self, doc: Document, template: Dict):
        """Add horizontal divider line"""
        primary = template["colors"]["primary"]
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), f"{primary[0]:02X}{primary[1]:02X}{primary[2]:02X}")
        pBdr.append(bottom)
        pPr.append(pBdr)
    
    def _add_section(self, doc: Document, title: str, template: Dict):
        """Add section header"""
        primary = template["colors"]["primary"]
        font_name = template["font"]
        
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(*primary)
        run.font.name = font_name
    
    def _add_paragraph(self, doc: Document, text: str, font_size: int = 10):
        """Add regular paragraph"""
        if not text:
            return
        para = doc.add_paragraph()
        run = para.add_run(str(text))
        run.font.size = Pt(font_size)
    
    def _add_experience(self, doc: Document, experience: List, template: Dict):
        """Add work experience section"""
        secondary = template["colors"]["secondary"]
        font_name = template["font"]
        
        for exp in experience:
            if not isinstance(exp, dict):
                continue
            
            # Company and title row
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(4)
            
            title_run = para.add_run(
                f"{exp.get('title', 'Position')} | {exp.get('company', 'Company')}"
            )
            title_run.bold = True
            title_run.font.size = Pt(10)
            title_run.font.name = font_name
            
            # Date
            if exp.get("duration") or exp.get("date"):
                date_para = doc.add_paragraph()
                date_run = date_para.add_run(
                    f"📅 {exp.get('duration') or exp.get('date', '')}"
                )
                date_run.font.size = Pt(9)
                date_run.font.color.rgb = RGBColor(*secondary)
                date_run.font.name = font_name
            
            # Description
            description = exp.get("description", "")
            if description:
                if isinstance(description, list):
                    for bullet in description:
                        bullet_para = doc.add_paragraph(style="List Bullet")
                        bullet_run = bullet_para.add_run(str(bullet))
                        bullet_run.font.size = Pt(9.5)
                        bullet_run.font.name = font_name
                else:
                    for line in str(description).split("\n"):
                        if line.strip():
                            bullet_para = doc.add_paragraph(style="List Bullet")
                            bullet_run = bullet_para.add_run(line.strip("• ").strip())
                            bullet_run.font.size = Pt(9.5)
                            bullet_run.font.name = font_name
    
    def _add_education(self, doc: Document, education: List, template: Dict):
        """Add education section"""
        secondary = template["colors"]["secondary"]
        font_name = template["font"]
        
        for edu in education:
            if not isinstance(edu, dict):
                continue
            
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(4)
            
            degree_text = f"{edu.get('degree', '')} | {edu.get('institution', edu.get('university', 'University'))}"
            degree_run = para.add_run(degree_text)
            degree_run.bold = True
            degree_run.font.size = Pt(10)
            degree_run.font.name = font_name
            
            details = []
            if edu.get("year") or edu.get("graduation_year"):
                details.append(f"📅 {edu.get('year') or edu.get('graduation_year')}")
            if edu.get("gpa"):
                details.append(f"GPA: {edu['gpa']}")
            
            if details:
                detail_para = doc.add_paragraph()
                detail_run = detail_para.add_run("  |  ".join(details))
                detail_run.font.size = Pt(9)
                detail_run.font.color.rgb = RGBColor(*secondary)
                detail_run.font.name = font_name
    
    def _add_skills(self, doc: Document, skills: List, template: Dict):
        """Add skills section"""
        font_name = template["font"]
        
        if not skills:
            return
        
        # Group skills by category
        categorized = {}
        for skill in skills:
            if isinstance(skill, dict):
                cat = skill.get("category", "Other")
                name = skill.get("name", str(skill))
            else:
                cat = "Technical Skills"
                name = str(skill)
            
            if cat not in categorized:
                categorized[cat] = []
            categorized[cat].append(name)
        
        # Display categorized
        for category, skill_list in categorized.items():
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(3)
            
            cat_run = para.add_run(f"{category}: ")
            cat_run.bold = True
            cat_run.font.size = Pt(9.5)
            cat_run.font.name = font_name
            
            skills_run = para.add_run(", ".join(skill_list))
            skills_run.font.size = Pt(9.5)
            skills_run.font.name = font_name
    
    def _add_projects(self, doc: Document, projects: List, template: Dict):
        """Add projects section"""
        secondary = template["colors"]["secondary"]
        font_name = template["font"]
        
        for project in projects:
            if not isinstance(project, dict):
                continue
            
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(4)
            
            name_run = para.add_run(project.get("name", "Project"))
            name_run.bold = True
            name_run.font.size = Pt(10)
            name_run.font.name = font_name
            
            if project.get("technologies"):
                tech = project["technologies"]
                if isinstance(tech, list):
                    tech = ", ".join(tech)
                tech_para = doc.add_paragraph()
                tech_run = tech_para.add_run(f"🛠️ {tech}")
                tech_run.font.size = Pt(9)
                tech_run.font.color.rgb = RGBColor(*secondary)
                tech_run.font.name = font_name
            
            if project.get("description"):
                desc_para = doc.add_paragraph(style="List Bullet")
                desc_run = desc_para.add_run(str(project["description"])[:200])
                desc_run.font.size = Pt(9.5)
                desc_run.font.name = font_name
    
    def _add_certifications(self, doc: Document, certifications: List, template: Dict):
        """Add certifications section"""
        font_name = template["font"]
        
        for cert in certifications:
            if not isinstance(cert, dict):
                para = doc.add_paragraph(style="List Bullet")
                run = para.add_run(str(cert))
                run.font.size = Pt(9.5)
                run.font.name = font_name
                continue
            
            para = doc.add_paragraph(style="List Bullet")
            cert_text = cert.get("name", "Certification")
            if cert.get("issuer"):
                cert_text += f" | {cert['issuer']}"
            if cert.get("year"):
                cert_text += f" ({cert['year']})"
            
            run = para.add_run(cert_text)
            run.font.size = Pt(9.5)
            run.font.name = font_name
