# app/services/resume_power_service.py
# 4 unique features: AI Tailoring, Interview Predictor, Resume Variants, Radar Score
import os
import json
from typing import Dict, List, Optional
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

load_dotenv()


class ResumePowerService:
    """4 unique AI-powered resume features that competitors don't have."""

    def __init__(self):
        api_key = os.getenv("GROQ_API_KEY")
        if not api_key:
            raise ValueError("GROQ_API_KEY not found")
        self.client = Groq(api_key=api_key)
        self.model = "llama-3.3-70b-versatile"
        self.output_dir = "generated_resumes"
        os.makedirs(self.output_dir, exist_ok=True)

    def _call_ai(self, system: str, user: str, max_tokens: int = 2000) -> str:
        resp = self.client.chat.completions.create(
            model=self.model,
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": user},
            ],
            temperature=0.7,
            max_tokens=max_tokens,
        )
        return resp.choices[0].message.content.strip()

    def _parse_json(self, raw: str) -> dict:
        try:
            start = raw.find("{")
            end = raw.rfind("}") + 1
            return json.loads(raw[start:end])
        except Exception:
            return {}

    # ─────────────────────────────────────────────────────────────────────────
    # FEATURE 1: AI RESUME TAILORING
    # Rewrites entire resume to match a specific job description
    # ─────────────────────────────────────────────────────────────────────────
    def tailor_resume(self, resume_data: dict, job_description: str, target_role: str) -> dict:
        """Rewrite resume content to perfectly match a job description."""

        system = """You are an elite resume writer who specializes in tailoring resumes 
        to specific job descriptions. You rewrite content to maximize keyword matches 
        and relevance without lying or fabricating experience."""

        user = f"""Tailor this resume for the following job description.
        
JOB DESCRIPTION:
{job_description}

TARGET ROLE: {target_role}

CURRENT RESUME DATA:
{json.dumps(resume_data, indent=2)}

Return ONLY a valid JSON object:
{{
  "tailoring_score": <number 0-100, how well original matched>,
  "improved_score": <number 0-100, how well tailored version matches>,
  "summary": "<a powerful 2-3 sentence professional summary tailored to this role>",
  "key_changes": ["<change 1>", "<change 2>", "<change 3>", "<change 4>", "<change 5>"],
  "tailored_experience": [
    {{
      "title": "<job title>",
      "company": "<company>",
      "duration": "<duration>",
      "description": "<rewritten description with job-relevant keywords and stronger impact verbs>"
    }}
  ],
  "highlighted_skills": ["<most relevant skill 1>", "<skill 2>", "<skill 3>", "<skill 4>", "<skill 5>", "<skill 6>"],
  "added_keywords": ["<keyword from JD added to resume>", "<keyword 2>", "<keyword 3>"],
  "cover_letter_opener": "<a powerful opening paragraph for a cover letter for this specific role>"
}}"""

        raw = self._call_ai(system, user, max_tokens=2500)
        result = self._parse_json(raw)

        if not result:
            return {"success": False, "error": "AI tailoring failed"}

        return {"success": True, "data": result}

    # ─────────────────────────────────────────────────────────────────────────
    # FEATURE 2: INTERVIEW QUESTIONS PREDICTOR
    # Predicts interview questions based on YOUR specific resume
    # ─────────────────────────────────────────────────────────────────────────
    def predict_interview_questions(self, resume_data: dict, target_role: str) -> dict:
        """Predict likely interview questions based on this specific resume."""

        system = """You are a senior hiring manager and interview coach. 
        You analyze a candidate's resume and predict the exact questions 
        an interviewer would ask based on their specific background, gaps, and achievements."""

        resume_summary = {
            "contact": resume_data.get("contact_info", {}),
            "experience": resume_data.get("experience", []),
            "skills": resume_data.get("skills", []),
            "education": resume_data.get("education", []),
        }

        user = f"""Based on this specific resume, predict the most likely interview questions 
a recruiter/hiring manager would ask this candidate for a {target_role} position.
Focus on their SPECIFIC background, not generic questions.

RESUME:
{json.dumps(resume_summary, indent=2)}

Return ONLY a valid JSON object:
{{
  "behavioral_questions": [
    {{
      "question": "<specific behavioral question based on their resume>",
      "why_asked": "<why interviewer would ask this>",
      "suggested_answer_structure": "<brief guidance on how to answer>"
    }}
  ],
  "technical_questions": [
    {{
      "question": "<technical question based on their specific skills>",
      "why_asked": "<based on which skill/experience>",
      "suggested_answer_structure": "<key points to cover>"
    }}
  ],
  "gap_questions": [
    {{
      "question": "<question about a gap or weakness in their resume>",
      "concern": "<what the interviewer is really asking>",
      "suggested_answer_structure": "<how to address this honestly>"
    }}
  ],
  "strength_questions": [
    {{
      "question": "<question highlighting their strongest point>",
      "opportunity": "<how to leverage this>",
      "suggested_answer_structure": "<key points>"
    }}
  ],
  "overall_interview_tips": ["<tip 1 specific to this candidate>", "<tip 2>", "<tip 3>"]
}}

Generate 3 questions per category."""

        raw = self._call_ai(system, user, max_tokens=3000)
        result = self._parse_json(raw)

        if not result:
            return {"success": False, "error": "Prediction failed"}

        return {"success": True, "data": result}

    # ─────────────────────────────────────────────────────────────────────────
    # FEATURE 3: RADAR SCORE — 6-dimension visual breakdown
    # ─────────────────────────────────────────────────────────────────────────
    def get_radar_score(self, resume_data: dict, target_role: str = None) -> dict:
        """Generate 6-dimension radar chart scores for visual display."""

        system = """You are a resume expert. Score resumes across 6 key dimensions 
        that recruiters evaluate. Be honest and strict — most resumes score 50-70."""

        user = f"""Score this resume across 6 dimensions for a {target_role or 'general'} role.

RESUME DATA:
{json.dumps(resume_data, indent=2)}

Return ONLY a valid JSON object with scores 0-100:
{{
  "dimensions": {{
    "impact": {{
      "score": <0-100>,
      "label": "Impact & Achievements",
      "feedback": "<1 sentence specific feedback>",
      "tip": "<1 actionable tip>"
    }},
    "clarity": {{
      "score": <0-100>,
      "label": "Clarity & Structure", 
      "feedback": "<1 sentence specific feedback>",
      "tip": "<1 actionable tip>"
    }},
    "ats": {{
      "score": <0-100>,
      "label": "ATS Compatibility",
      "feedback": "<1 sentence specific feedback>",
      "tip": "<1 actionable tip>"
    }},
    "keywords": {{
      "score": <0-100>,
      "label": "Keyword Density",
      "feedback": "<1 sentence specific feedback>",
      "tip": "<1 actionable tip>"
    }},
    "experience": {{
      "score": <0-100>,
      "label": "Experience Quality",
      "feedback": "<1 sentence specific feedback>",
      "tip": "<1 actionable tip>"
    }},
    "skills": {{
      "score": <0-100>,
      "label": "Skills Relevance",
      "feedback": "<1 sentence specific feedback>",
      "tip": "<1 actionable tip>"
    }}
  }},
  "overall": <average of all 6 scores>,
  "strongest_dimension": "<name of highest scoring dimension>",
  "weakest_dimension": "<name of lowest scoring dimension>",
  "priority_action": "<single most important thing to improve>"
}}"""

        raw = self._call_ai(system, user, max_tokens=1500)
        result = self._parse_json(raw)

        if not result:
            return {"success": False, "error": "Scoring failed"}

        return {"success": True, "data": result}

    # ─────────────────────────────────────────────────────────────────────────
    # FEATURE 4: ONE-CLICK RESUME VARIANTS
    # Generate 3 tone variants as DOCX files simultaneously
    # ─────────────────────────────────────────────────────────────────────────
    def generate_variants(self, resume_data: dict, user_id: int) -> dict:
        """Generate 3 resume tone variants: Aggressive, Conservative, Technical."""

        system = """You are an expert resume writer. Rewrite resume content in 3 
        distinct tones/styles. Keep facts accurate but change emphasis and language."""

        user = f"""Rewrite this resume content in 3 distinct variants.

RESUME DATA:
{json.dumps(resume_data, indent=2)}

Return ONLY a valid JSON object:
{{
  "aggressive": {{
    "tone_description": "Bold, achievement-focused, strong action verbs, quantified results",
    "summary": "<powerful aggressive summary>",
    "experience_bullets": [
      "<rewritten bullet - bold and achievement focused>",
      "<bullet 2>",
      "<bullet 3>",
      "<bullet 4>",
      "<bullet 5>"
    ]
  }},
  "conservative": {{
    "tone_description": "Formal, traditional, team-oriented, responsibility-focused",
    "summary": "<professional conservative summary>",
    "experience_bullets": [
      "<rewritten bullet - formal and responsibility focused>",
      "<bullet 2>",
      "<bullet 3>",
      "<bullet 4>",
      "<bullet 5>"
    ]
  }},
  "technical": {{
    "tone_description": "Skill-heavy, technology-focused, specific tools and methodologies",
    "summary": "<technical skill-focused summary>",
    "experience_bullets": [
      "<rewritten bullet - emphasizes technical skills and tools>",
      "<bullet 2>",
      "<bullet 3>",
      "<bullet 4>",
      "<bullet 5>"
    ]
  }}
}}"""

        raw = self._call_ai(system, user, max_tokens=2000)
        variants_data = self._parse_json(raw)

        if not variants_data:
            return {"success": False, "error": "Variant generation failed"}

        # Generate a DOCX for each variant
        files = {}
        contact = resume_data.get("contact_info", {})
        name = contact.get("name", "Candidate")

        for variant_key in ["aggressive", "conservative", "technical"]:
            variant = variants_data.get(variant_key, {})
            if not variant:
                continue

            doc = Document()

            # Header
            h = doc.add_heading(name, 0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in h.runs:
                run.font.color.rgb = RGBColor(
                    *{"aggressive": (192, 0, 0), "conservative": (31, 73, 125), "technical": (0, 112, 192)}[variant_key]
                )

            # Contact line
            contact_parts = [x for x in [
                contact.get("email"), contact.get("phone"),
                contact.get("linkedin"), contact.get("location")
            ] if x]
            if contact_parts:
                p = doc.add_paragraph(" | ".join(contact_parts))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)

            doc.add_paragraph()

            # Tone badge
            tone_p = doc.add_paragraph()
            tone_run = tone_p.add_run(f"[{variant_key.upper()} VARIANT] — {variant.get('tone_description', '')}")
            tone_run.font.size = Pt(8)
            tone_run.font.italic = True
            tone_run.font.color.rgb = RGBColor(128, 128, 128)

            # Summary
            doc.add_heading("Professional Summary", 1)
            doc.add_paragraph(variant.get("summary", ""))

            # Experience bullets
            if variant.get("experience_bullets"):
                doc.add_heading("Key Achievements", 1)
                for bullet in variant["experience_bullets"]:
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(bullet)

            # Original experience
            experience = resume_data.get("experience", [])
            if experience:
                doc.add_heading("Experience", 1)
                for exp in experience[:4]:
                    title = exp.get("title", "") if isinstance(exp, dict) else str(exp)
                    company = exp.get("company", "") if isinstance(exp, dict) else ""
                    duration = exp.get("duration", "") if isinstance(exp, dict) else ""
                    p = doc.add_paragraph()
                    run = p.add_run(f"{title} — {company}")
                    run.bold = True
                    if duration:
                        p.add_run(f"  |  {duration}").font.size = Pt(9)

            # Skills
            skills = resume_data.get("skills", [])
            if skills:
                doc.add_heading("Skills", 1)
                skill_names = [
                    (s.get("name") if isinstance(s, dict) else str(s))
                    for s in skills[:15]
                ]
                doc.add_paragraph(" • ".join(filter(None, skill_names)))

            # Education
            education = resume_data.get("education", [])
            if education:
                doc.add_heading("Education", 1)
                for edu in education[:2]:
                    if isinstance(edu, dict):
                        degree = edu.get("degree", "")
                        school = edu.get("school", "")
                        year = edu.get("year", "")
                        p = doc.add_paragraph()
                        p.add_run(f"{degree} — {school}").bold = True
                        if year:
                            p.add_run(f"  ({year})")

            filepath = os.path.join(self.output_dir, f"resume_{variant_key}_{user_id}.docx")
            doc.save(filepath)
            files[variant_key] = filepath

        return {
            "success": True,
            "variants_data": variants_data,
            "files": {k: os.path.basename(v) for k, v in files.items()},
            "file_paths": files,
        }